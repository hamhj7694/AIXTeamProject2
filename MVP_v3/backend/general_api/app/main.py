from __future__ import annotations

import asyncio
import io
import os
import hashlib
import logging
import mimetypes
import re
import secrets
from datetime import datetime, timedelta, timezone
from uuid import uuid4
from pathlib import Path
from typing import Literal

from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, Request
from fastapi.exception_handlers import request_validation_exception_handler
from fastapi.exceptions import RequestValidationError
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field, model_validator
from .domains.cases.context_items import Section, ContextItemChange, ContextItemConflictError
from .domains.cases.context_item_repository import ContextItemRepository, InMemoryContextItemRepository

from contracts.diagnosis import AnalyzeTextRequest
from contracts.public_api.customer_progress import CustomerProgressItem, ProgressStep, UpdateCustomerProgress
from .domains.cases.customer_progress import PREFIX as PROGRESS_PREFIX, ProgressConflict, progress_items as build_customer_progress, progress_ai_context, actions_for_ai
from request_trace import install_request_trace
from contracts.ai_internal.work_card import CaseWorkCardOutput, WorkCardType
from contracts.public_api.case_analyze import (
    PublicAnalyzeCaseRequest,
    PublicAnalyzeCaseResponse,
    PublicAnalyzeError,
    PublicInitialReportReference,
)
from contracts.public_api.case_read import PublicCaseReadResponse, to_public_case_read_response, to_public_case_summary_response
from contracts.public_api.case_transition import PublicCasePatchRequest
from contracts.public_api.case_activity import (
    PublicCaseEventResponse,
    PublicAttachmentResponse,
    PublicCreateMessageRequest,
    PublicCustomerEmergencyRequest,
    PublicMessageResponse,
    to_public_event,
    to_public_attachment,
    to_public_message,
)
from contracts.public_api.case_workflow import (
    PublicActionResponse,
    PublicAnswerCustomerQuestionRequest,
    PublicCaseFactResponse,
    PublicConfirmCaseFactRequest,
    PublicPersonalNoteCreateRequest,
    PublicPersonalNoteUpdateRequest,
    PublicPersonalNoteResponse,
    PublicCaseBundleResponse,
    PublicCreateActionRequest,
    PublicUpdateActionRequest,
    PublicActionCommandRequest,
    PublicCreateVoiceSessionRequest,
    PublicUpdateVoiceSessionRequest,
    PublicVoiceSessionResponse,
    PublicCreateTranscriptRequest,
    PublicTranscriptResponse,
    PublicReportResponse,
    PublicCreateVerificationRequest,
    PublicCustomerQuestionResponse,
    PublicCustomerQuestionView,
    PublicCustomerVerificationResult,
    PublicQuestionCandidateResponse,
    PublicCaseContextProjection,
    PublicCaseSupportBrief,
    PublicCaseSupportSnapshotResponse,
    PublicUnresolvedItemResponse,
    PublicQueueCustomerQuestionsRequest,
    PublicUpdateVerificationRequest,
    PublicVerificationResponse,
    to_public_customer_question,
    to_public_customer_question_view,
    to_public_action,
    to_public_verification,
)
from contracts.public_api.collaboration import (
    MessageChannel,
    PublicAiInvocationRequest,
    PublicAiInvocationResponse,
    PublicAiShareRequest,
    PublicCustomerAiReplyRequest,
    PublicCaseMemberResponse,
    PublicCaseMemberUpsertRequest,
    PublicCasePresenceResponse,
    PublicPresenceHeartbeatRequest,
    PublicPrimaryAssigneeRequest,
    PublicPrimaryAssigneeResponse,
)

from .clients.diagnosis_ai import AiServiceAuthenticationError, AiServiceError, AiServiceQuotaError, HttpDiagnosisAiClient
from .domains.cases.repository import CasePersistenceError, CaseVersionConflictError, normalize_target_field
from .domains.cases.context_projection_repository import ContextProjectionRepository
from .domains.cases.mysql_repository import MySqlCaseRepository
from .domains.cases.service import AnalyzeCaseService, InvalidCaseTransitionError, transition_case


load_dotenv(Path(__file__).resolve().parents[3] / ".env", override=False)


def build_repository():
    repository_type = os.getenv("CASE_REPOSITORY", "mysql").lower()
    if repository_type == "mysql":
        return MySqlCaseRepository()
    raise RuntimeError(f"Unsupported CASE_REPOSITORY: {repository_type}. Use mysql in deployed environments.")

app = FastAPI(title="AI Independent Verification - General API", version="0.1.0")
install_request_trace(app, "general-api")
logger = logging.getLogger(__name__)

AUTONOMOUS_P0_QUESTION_FIELDS = {
    "transfer_status",
    "personal_information_exposure",
    "authentication_information_exposure",
    "remote_control_app",
}
AI_CHECKLIST_ACTION_PREFIX = "AI_CHECKLIST:"
STAFF_JUDGMENT_ACTION_TYPE = "STAFF_JUDGMENT"
CHECKLIST_FIELD_LABELS = {
    "transfer_status": "실제 송금 여부",
    "transfer_purpose": "송금 목적",
    "claimed_organization": "사칭 기관",
    "incident_claim": "상대방 주장",
    "personal_information_exposure": "개인정보 제공 여부",
    "authentication_information_exposure": "인증정보 제공 여부",
    "remote_control_app": "원격제어 앱 설치 여부",
}
PROACTIVE_CASE_POLL_SECONDS = max(1.0, float(os.getenv("PROACTIVE_CASE_POLL_SECONDS", "3")))
_proactive_case_revisions: dict[str, str] = {}
_proactive_worker_task: asyncio.Task | None = None


class AdminCaseDeleteRequest(BaseModel):
    password: str


class AdminCaseFinalizeRequest(BaseModel):
    expected_version: int = Field(ge=1)
    password: str
    note: str = Field(default="", max_length=10_000)


class AdminCaseReopenRequest(BaseModel):
    expected_version: int = Field(ge=1)
    password: str


class CaseOutcomeRequest(BaseModel):
    expected_version: int
    victim_transfer_status: Literal["UNKNOWN", "YES", "NO"]
    actual_loss_amount_krw: float | None = None


class PublicWorkCardGenerateRequest(BaseModel):
    card_type: WorkCardType
default_cors_origins = (
    "http://localhost:5173,http://127.0.0.1:5173,"
    "http://localhost:5174,http://127.0.0.1:5174,"
    "http://localhost:5175,http://127.0.0.1:5175,"
    "http://localhost:5176,http://127.0.0.1:5176"
)
cors_origins = [
    origin.strip()
    for origin in os.getenv("CORS_ALLOWED_ORIGINS", default_cors_origins).split(",")
    if origin.strip()
]
app.add_middleware(
    CORSMiddleware,
    allow_origins=cors_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
repository = build_repository()

ATTACHMENT_STORAGE_ROOT = Path(os.getenv("ATTACHMENT_STORAGE_ROOT", str(Path(__file__).resolve().parents[2] / "data" / "uploads"))).resolve()
MAX_ATTACHMENT_BYTES = int(os.getenv("MAX_ATTACHMENT_BYTES", str(10 * 1024 * 1024)))
ALLOWED_ATTACHMENT_TYPES = {
    "image/jpeg": ".jpg", "image/png": ".png", "image/gif": ".gif", "image/webp": ".webp",
    "application/pdf": ".pdf", "text/plain": ".txt", "text/csv": ".csv", "application/json": ".json",
    "application/msword": ".doc", "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "application/vnd.ms-excel": ".xls", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
}


def _clean_file_name(value: str) -> str:
    name = Path(value).name
    name = re.sub(r"[\x00-\x1f\x7f]", "", name).strip()
    return name[:160] or "attachment"


def _resolve_attachment_type(file_name: str, content_type: str | None) -> tuple[str, str]:
    candidate = (content_type or "").split(";", 1)[0].strip().lower()
    if candidate not in ALLOWED_ATTACHMENT_TYPES:
        candidate = (mimetypes.guess_type(file_name)[0] or "").lower()
    extension = ALLOWED_ATTACHMENT_TYPES.get(candidate)
    if not extension:
        raise HTTPException(status_code=415, detail={"code": "ATTACHMENT_TYPE_NOT_ALLOWED", "message": "지원하지 않는 파일 형식입니다."})
    return candidate, extension


def _has_valid_signature(mime_type: str, content: bytes) -> bool:
    signatures = {
        "image/jpeg": lambda value: value.startswith(b"\xff\xd8\xff"),
        "image/png": lambda value: value.startswith(b"\x89PNG\r\n\x1a\n"),
        "image/gif": lambda value: value.startswith((b"GIF87a", b"GIF89a")),
        "image/webp": lambda value: len(value) >= 12 and value[:4] == b"RIFF" and value[8:12] == b"WEBP",
        "application/pdf": lambda value: value.startswith(b"%PDF-"),
        "application/msword": lambda value: value.startswith(b"\xd0\xcf\x11\xe0"),
        "application/vnd.ms-excel": lambda value: value.startswith(b"\xd0\xcf\x11\xe0"),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": lambda value: value.startswith(b"PK\x03\x04"),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": lambda value: value.startswith(b"PK\x03\x04"),
    }
    validator = signatures.get(mime_type)
    if validator:
        return validator(content)
    if mime_type.startswith("text/") or mime_type == "application/json":
        try:
            content.decode("utf-8")
            return True
        except UnicodeDecodeError:
            return False
    return True
service = AnalyzeCaseService(HttpDiagnosisAiClient(), repository)


@app.middleware("http")
async def request_id_middleware(request: Request, call_next):
    request_id = request.headers.get("X-Request-ID") or f"req-{uuid4().hex}"
    response = await call_next(request)
    response.headers["X-Request-ID"] = request_id
    return response


def public_failed_response(code: str, message: str, *, retryable: bool) -> PublicAnalyzeCaseResponse:
    return PublicAnalyzeCaseResponse(
        disposition="FAILED",
        error=PublicAnalyzeError(code=code, message=message, retryable=retryable),
    )


def to_public_analyze_response(result) -> PublicAnalyzeCaseResponse:
    if result.disposition == "FAILED":
        error = result.error
        return public_failed_response(
            error.code if error else "AI_ANALYSIS_FAILED",
            error.message if error else "진단을 완료하지 못했습니다.",
            retryable=error.retryable if error else True,
        )

    if result.disposition == "NO_CASE":
        return PublicAnalyzeCaseResponse(
            disposition="NO_CASE",
            risk="NORMAL",
            initial_brief=result.initial_brief,
        )

    report = result.initial_report
    return PublicAnalyzeCaseResponse(
        disposition="CASE_CREATED",
        case_id=result.case_id,
        risk=result.risk.value if hasattr(result.risk, "value") else result.risk,
        mode=result.mode,
        status=result.status,
        initial_brief=result.initial_brief,
        initial_report=PublicInitialReportReference(
            report_id=report.report_id,
            case_id=report.case_id,
            report_version=report.report_version,
        ) if report else None,
    )


@app.get("/")
async def root() -> dict[str, str]:
    return {"service": "general-api", "status": "ok", "health": "/health"}


@app.get("/health")
async def health() -> dict[str, str]:
    repository_type = os.getenv("CASE_REPOSITORY", "mysql").lower()
    ping = getattr(repository, "ping", None)
    if callable(ping):
        try:
            await ping()
        except Exception as error:
            logger.exception("Database health check failed")
            raise HTTPException(
                status_code=503,
                detail={"code": "DATABASE_UNAVAILABLE", "message": "데이터베이스에 연결할 수 없습니다."},
            ) from error
    return {"status": "ok", "database": repository_type}


@app.exception_handler(RequestValidationError)
async def public_analyze_validation_error(request: Request, exc: RequestValidationError):
    if request.url.path == "/api/cases/analyze":
        failure = public_failed_response("INVALID_INPUT", "요청 형식을 확인해 주세요.", retryable=False)
        return JSONResponse(status_code=400, content=failure.model_dump(mode="json"))
    return await request_validation_exception_handler(request, exc)


@app.post("/api/cases/analyze", response_model=PublicAnalyzeCaseResponse, status_code=201)
async def analyze_case(request: PublicAnalyzeCaseRequest) -> PublicAnalyzeCaseResponse | JSONResponse:
    internal_request = AnalyzeTextRequest.model_validate(request.model_dump())
    try:
        return to_public_analyze_response(await service.analyze(internal_request))
    except ValueError as exc:
        failure = public_failed_response("INVALID_INPUT", str(exc), retryable=False)
        return JSONResponse(status_code=400, content=failure.model_dump(mode="json"))
    except AiServiceQuotaError as exc:
        failure = public_failed_response("OPENAI_QUOTA_EXHAUSTED", str(exc), retryable=False)
        return JSONResponse(status_code=429, content=failure.model_dump(mode="json"))
    except AiServiceAuthenticationError as exc:
        failure = public_failed_response("OPENAI_AUTHENTICATION_FAILED", str(exc), retryable=False)
        return JSONResponse(status_code=401, content=failure.model_dump(mode="json"))
    except AiServiceError as exc:
        failure = public_failed_response("AI_ANALYSIS_FAILED", str(exc), retryable=True)
        return JSONResponse(status_code=503, content=failure.model_dump(mode="json"))
    except CasePersistenceError:
        failure = public_failed_response("CASE_SAVE_FAILED", "AI 분석 후 사건 저장을 완료하지 못했습니다.", retryable=True)
        return JSONResponse(status_code=503, content=failure.model_dump(mode="json"))
    except Exception as exc:
        failure = public_failed_response("AI_ANALYSIS_FAILED", "진단을 완료하지 못했습니다.", retryable=True)
        return JSONResponse(status_code=503, content=failure.model_dump(mode="json"))


async def to_case_read(record: dict) -> PublicCaseReadResponse:
    members = await repository.list_members(record["case_id"])
    primary = next((item.get("display_name") for item in members if item.get("role") == "CASE_OWNER"), None)
    deleted_at = record.get("deleted_at")
    trash_expires_at = None
    if deleted_at:
        deleted_instant = datetime.fromisoformat(str(deleted_at).replace("Z", "+00:00"))
        if deleted_instant.tzinfo is None:
            deleted_instant = deleted_instant.replace(tzinfo=timezone.utc)
        trash_expires_at = (deleted_instant + timedelta(days=30)).isoformat()
    return to_public_case_read_response({
        **record,
        "primary_assignee": primary,
        "trash_expires_at": trash_expires_at,
    })


@app.get("/api/cases", response_model=list[PublicCaseReadResponse], response_model_exclude_none=True)
async def list_cases() -> list[PublicCaseReadResponse]:
    return [await to_case_read(record) for record in await repository.list()]


def require_admin_password(password: str) -> None:
    configured_password = os.getenv("CASE_ADMIN_DELETE_PASSWORD")
    if not configured_password:
        raise HTTPException(status_code=503, detail={"code": "ADMIN_AUTH_NOT_CONFIGURED", "message": "관리자 인증 설정이 필요합니다."})
    if not secrets.compare_digest(password, configured_password):
        raise HTTPException(status_code=403, detail={"code": "ADMIN_AUTH_FAILED", "message": "관리자 비밀번호가 올바르지 않습니다."})


@app.post("/api/cases/admin/verify-password", status_code=204)
async def verify_admin_password(request: AdminCaseDeleteRequest) -> None:
    require_admin_password(request.password)


@app.get("/api/cases/trash", response_model=list[PublicCaseReadResponse], response_model_exclude_none=True)
async def list_trashed_cases() -> list[PublicCaseReadResponse]:
    return [await to_case_read(record) for record in await repository.list_trashed_cases()]


@app.post("/api/cases/{case_id}/trash", status_code=204)
async def move_case_to_trash(case_id: str, request: AdminCaseDeleteRequest) -> None:
    """Keep a Case in trash for 30 days after administrator verification."""
    require_admin_password(request.password)
    try:
        await repository.delete_case(case_id)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case not found."}) from exc


@app.delete("/api/cases/{case_id}", status_code=204)
async def move_case_to_trash_legacy(case_id: str, request: AdminCaseDeleteRequest) -> None:
    """Backward-compatible alias for moving a Case to trash."""
    await move_case_to_trash(case_id, request)


@app.post("/api/cases/{case_id}/restore", status_code=204)
async def restore_case_from_trash(case_id: str, request: AdminCaseDeleteRequest) -> None:
    require_admin_password(request.password)
    try:
        await repository.restore_case(case_id)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case not found in trash."}) from exc


@app.delete("/api/cases/trash/{case_id}", status_code=204)
async def permanently_delete_trashed_case(case_id: str, request: AdminCaseDeleteRequest) -> None:
    require_admin_password(request.password)
    attachments = await repository.list_attachments(case_id)
    try:
        await repository.purge_case(case_id)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case not found in trash."}) from exc
    for attachment in attachments:
        stored_path = (ATTACHMENT_STORAGE_ROOT / attachment["storage_path"]).resolve()
        if ATTACHMENT_STORAGE_ROOT not in stored_path.parents:
            logger.error("Skipped unsafe attachment path during Case purge: %s", stored_path)
            continue
        try:
            stored_path.unlink(missing_ok=True)
        except OSError:
            logger.exception("Failed to remove attachment after Case purge: %s", stored_path)


@app.get("/api/cases/{case_id}", response_model=PublicCaseReadResponse, response_model_exclude_none=True)
async def get_case(case_id: str) -> PublicCaseReadResponse:
    record = await repository.get(case_id)
    if record is None:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case를 찾을 수 없습니다."})
    return await to_case_read(record)


@app.patch("/api/cases/{case_id}", response_model=PublicCaseReadResponse, response_model_exclude_none=True)
async def patch_case(case_id: str, request: PublicCasePatchRequest) -> PublicCaseReadResponse:
    try:
        record = await transition_case(
            repository,
            case_id,
            request.expected_version,
            status=request.status,
            mode=request.mode,
        )
    except KeyError as exc:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case not found."}) from exc
    except CaseVersionConflictError as exc:
        raise HTTPException(
            status_code=409,
            detail={"code": "VERSION_CONFLICT", "message": "Case has changed.", "current_version": exc.current_version},
        ) from exc
    except InvalidCaseTransitionError as exc:
        raise HTTPException(status_code=409, detail={"code": "INVALID_STATE_TRANSITION", "message": str(exc)}) from exc
    return await to_case_read(record)


@app.put("/api/cases/{case_id}/outcome", response_model=PublicCaseReadResponse, response_model_exclude_none=True)
async def update_case_outcome(case_id: str, request: CaseOutcomeRequest) -> PublicCaseReadResponse:
    try:
        record = await repository.update_case(case_id, request.expected_version, {
            "victim_transfer_status": request.victim_transfer_status,
            "actual_loss_amount_krw": request.actual_loss_amount_krw if request.victim_transfer_status == "YES" else None,
        })
    except KeyError as exc:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case not found."}) from exc
    except CaseVersionConflictError as exc:
        raise HTTPException(status_code=409, detail={"code": "VERSION_CONFLICT", "message": "Case has changed.", "current_version": exc.current_version}) from exc
    return await to_case_read(record)


async def require_case(case_id: str) -> None:
    if await repository.get(case_id) is None:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case를 찾을 수 없습니다."})


def build_customer_question_candidates(case: dict, queued: list[dict]) -> list[PublicQuestionCandidateResponse]:
    """Deterministic MVP candidates. AI may replace this source, not the queue contract."""
    already_handled = {item["target_field"] for item in queued if item.get("status") in {"PENDING", "ASKED", "ANSWERED"}}
    fields = [
        ("victim_transfer_status", "현재 송금하거나 이체한 금액이 있나요?", "피해 여부와 피해 금액을 먼저 확인해야 합니다.", "안전을 위해 피해 발생 여부를 가장 먼저 확인합니다.", "P0", ["없음", "있음", "잘 모르겠어요"]),
        ("remote_control_app", "휴대폰에 원격 제어 또는 화면 공유 앱을 설치하라는 안내를 받으셨나요?", "추가 피해 가능성을 확인해야 합니다.", "휴대폰 제어 가능성을 확인해 추가 피해를 막기 위한 질문입니다.", "P0", ["설치함", "설치하지 않음", "잘 모르겠어요"]),
        ("personal_information_exposure", "주민등록번호나 계좌번호 등 개인정보를 제공하셨나요?", "개인정보 노출 여부는 추가 보호 조치 판단에 필요합니다.", "개인정보 보호 조치가 필요한지 확인하는 질문입니다.", "P0", ["제공하지 않았어요", "일부 제공했어요", "제공했어요", "잘 모르겠어요"]),
        ("authentication_information_exposure", "인증번호, 비밀번호 또는 OTP를 제공하셨나요?", "인증정보 노출 여부는 계정 보호 판단에 필요합니다.", "계정과 금융정보를 보호하기 위해 인증정보 노출 여부를 확인합니다.", "P0", ["제공하지 않았어요", "제공했어요", "잘 모르겠어요"]),
        ("impersonated_institution", "상대방이 어느 기관이나 은행을 사칭했는지 알려주실 수 있나요?", "공식 채널 검증 대상을 정해야 합니다.", "상대방의 주장을 공식 채널에서 확인하기 위한 질문입니다.", "P1", []),
    ]
    if case.get("victim_transfer_status") != "UNKNOWN":
        already_handled.add("victim_transfer_status")
    return [
        PublicQuestionCandidateResponse(
            question_id=f"candidate-{target_field}", target_field=target_field,
            question_text=text, reason=reason, customer_explanation=customer_explanation,
            priority=priority, options=options, answer_mode="CHOICE_OR_TEXT", allow_free_text=True,
        )
        for target_field, text, reason, customer_explanation, priority, options in fields
        if target_field not in already_handled
    ]


def build_question_recommendation_context(facts: list[dict], questions: list[dict], case: dict | None = None) -> dict:
    """답변 수신과 사실 확정을 분리하되 질문 이력이 있는 항목은 다시 묻지 않는다."""
    valid_fields = {
        "transfer_status", "transfer_purpose", "claimed_organization", "incident_claim",
        "personal_information_exposure", "authentication_information_exposure",
        "remote_control_app",
    }
    confirmed_fields = [normalize_target_field(item["field"]) for item in facts if item.get("status") == "CONFIRMED" and normalize_target_field(item.get("field", "")) in valid_fields]
    if case and case.get("victim_transfer_status") in {"YES", "NO"}:
        confirmed_fields.append("transfer_status")
    return {
        "confirmed_fields": list(dict.fromkeys(confirmed_fields)),
        "pending_question_fields": [normalize_target_field(item["target_field"]) for item in questions if item.get("status") in {"PENDING", "ASKED"} and normalize_target_field(item.get("target_field", "")) in valid_fields],
        "answered_question_fields": [normalize_target_field(item["target_field"]) for item in questions if item.get("status") == "ANSWERED" and normalize_target_field(item.get("target_field", "")) in valid_fields],
        "answered_question_ids": [item["question_id"] for item in questions if item.get("status") == "ANSWERED"],
    }


def normalize_question_text(value: str) -> str:
    """Whitespace/case differences must not bypass the duplicate-question guard."""
    return " ".join(value.split()).casefold()


def exclude_handled_question_candidates(
    candidates: list[PublicQuestionCandidateResponse], questions: list[dict]
) -> list[PublicQuestionCandidateResponse]:
    handled = [item for item in questions if item.get("status") in {"PENDING", "ASKED", "ANSWERED"}]
    handled_fields = {normalize_target_field(str(item.get("target_field", ""))) for item in handled}
    handled_texts = {normalize_question_text(str(item.get("question_text", ""))) for item in handled}
    return [
        candidate for candidate in candidates
        if normalize_target_field(candidate.target_field) not in handled_fields
        and normalize_question_text(candidate.question_text) not in handled_texts
    ]


def to_public_case_support_snapshot(
    case_id: str, payload: dict, *, available: bool, source_revision: int | None = None,
    projection_revision: int | None = None, projection_status: str = "UNCACHED",
) -> PublicCaseSupportSnapshotResponse:
    brief = payload.get("case_brief") or None
    context = payload.get("case_context") or None
    return PublicCaseSupportSnapshotResponse(
        case_id=case_id, available=available,
        case_brief=PublicCaseSupportBrief(
            summary=brief["summary"], incident_type=brief["incident_type"],
            risk_level=brief["risk_level"], risk_score=brief["risk_score"], next_checks=brief.get("next_checks", []),
        ) if brief else None,
        case_context=PublicCaseContextProjection.model_validate(context) if context else None,
        recommended_questions=[PublicQuestionCandidateResponse(
            question_id=item["question_id"], target_field=item["target_field"],
            question_text=item["question"], reason=item["reason"], priority=item["priority"],
            options={
                "transfer_status": ["아직 송금하지 않았어요", "이미 송금했어요", "잘 모르겠어요"],
                "personal_information_exposure": ["제공하지 않았어요", "일부 제공했어요", "제공했어요", "잘 모르겠어요"],
                "authentication_information_exposure": ["제공하지 않았어요", "제공했어요", "잘 모르겠어요"],
            }.get(item["target_field"], []),
            customer_explanation={
                "transfer_status": "추가 피해를 막고 필요한 조치를 안내하기 위해 실제 송금 여부를 먼저 확인합니다.",
                "personal_information_exposure": "개인정보 보호 조치가 필요한지 확인하는 질문입니다.",
                "authentication_information_exposure": "계정과 금융정보를 보호하기 위해 인증정보 노출 여부를 확인합니다.",
            }.get(item["target_field"]),
        ) for item in payload.get("recommended_questions", [])],
        unresolved_items=[PublicUnresolvedItemResponse.model_validate(item) for item in payload.get("unresolved_items", [])],
        warnings=payload.get("warnings", []),
        source_revision=source_revision, projection_revision=projection_revision,
        projection_status=projection_status,
    )


async def _read_case_support_source(case_id: str, *, attempts: int = 3) -> tuple[int, dict, list[dict], list[dict], list[dict], list[dict]]:
    """Read a source set whose semantic revision did not change mid-read."""
    for _ in range(attempts):
        case = await repository.get(case_id)
        if case is None:
            raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case를 찾을 수 없습니다."})
        before = int(case.get("context_revision", 1))
        facts, questions, verifications, actions = await asyncio.gather(
            repository.list_case_facts(case_id), repository.list_customer_questions(case_id),
            repository.list_verifications(case_id), repository.list_actions(case_id),
        )
        latest = await repository.get(case_id)
        if latest is None:
            raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case를 찾을 수 없습니다."})
        after = int(latest.get("context_revision", 1))
        if before == after:
            return after, latest, facts, questions, verifications, actions
    raise RuntimeError("CASE_CONTEXT_SOURCE_CHANGED")


def _case_support_ai_input(case_id: str, case: dict, facts: list[dict], questions: list[dict], verifications: list[dict], actions: list[dict]) -> dict:
    actions = actions_for_ai(actions)
    return {
        "case_id": case_id, "diagnosis": case.get("diagnosis"),
        "question_context": build_question_recommendation_context(facts, questions, case),
        "questions": [{
            "question_id": item["question_id"],
            "target_field": normalize_target_field(str(item.get("target_field", ""))),
            "question_text": item.get("question_text", ""), "priority": item.get("priority", "P1"),
            "status": item.get("status", "PENDING"), "answer_text": item.get("answer_text"),
        } for item in questions],
        "facts": [{
            "fact_id": item["fact_id"], "field": normalize_target_field(str(item.get("field", item.get("field_name", "")))),
            "value": str(item.get("value", "")), "status": item.get("status", "UNRESOLVED"),
        } for item in facts],
        "verifications": [{
            "verification_task_id": item["verification_task_id"], "target": item.get("target", ""),
            "claim": item.get("claim", ""), "status": item.get("status", "PENDING"),
            "result_summary": item.get("result_summary"),
        } for item in verifications],
        "actions": [{
            "action_id": item["action_id"], "action_type": item.get("action_type", "OTHER"),
            "status": item.get("status", "PENDING"), "note": item.get("note", ""),
        } for item in actions],
    }


async def get_case_support_snapshot(case_id: str) -> PublicCaseSupportSnapshotResponse:
    # Mock/in-memory repositories keep the original uncached path. Production
    # MySQL uses durable revision + DB lease so multiple workers share one result.
    if isinstance(repository, MySqlCaseRepository):
        projections = ContextProjectionRepository(repository)
        for _ in range(3):
            try:
                revision, case, facts, questions, verifications, actions = await _read_case_support_source(case_id)
            except RuntimeError:
                continue
            claim = await projections.claim(case_id, revision)
            if claim.outcome == "STALE":
                continue
            if claim.outcome == "CACHED":
                return to_public_case_support_snapshot(case_id, claim.last_success_payload or {}, available=True,
                    source_revision=revision, projection_revision=claim.last_success_revision, projection_status="CURRENT")
            if claim.outcome == "IN_PROGRESS":
                if claim.last_success_payload is not None:
                    cached = {**claim.last_success_payload, "warnings": [
                        *claim.last_success_payload.get("warnings", []), "최신 변경사항을 반영 중이며 직전 정상 사건 맥락을 표시합니다.",
                    ]}
                    return to_public_case_support_snapshot(case_id, cached, available=True,
                        source_revision=revision, projection_revision=claim.last_success_revision, projection_status="UPDATING")
                return PublicCaseSupportSnapshotResponse(case_id=case_id, available=False,
                    warnings=["사건 맥락을 처음 정리하고 있습니다."], source_revision=revision, projection_status="UPDATING")
            try:
                payload = await service.ai_client.build_case_support_snapshot(
                    _case_support_ai_input(case_id, case, facts, questions, verifications, actions)
                )
            except AiServiceError as exc:
                await projections.fail(case_id, revision, claim.lease_token or "", type(exc).__name__)
                if claim.last_success_payload is not None:
                    cached = {**claim.last_success_payload, "warnings": [
                        *claim.last_success_payload.get("warnings", []), str(exc), "최신 반영에 실패해 직전 정상 사건 맥락을 표시합니다.",
                    ]}
                    return to_public_case_support_snapshot(case_id, cached, available=True,
                        source_revision=revision, projection_revision=claim.last_success_revision, projection_status="STALE")
                return PublicCaseSupportSnapshotResponse(case_id=case_id, available=False, warnings=[str(exc)],
                    source_revision=revision, projection_status="FAILED")
            if await projections.complete(case_id, revision, claim.lease_token or "", payload):
                return to_public_case_support_snapshot(case_id, payload, available=True,
                    source_revision=revision, projection_revision=revision, projection_status="CURRENT")
            # Data changed during generation; never publish this obsolete result.
        return PublicCaseSupportSnapshotResponse(case_id=case_id, available=False,
            warnings=["사건 정보가 연속으로 변경되어 최신 맥락 반영을 다시 시도합니다."], projection_status="UPDATING")

    case = await repository.get(case_id)
    if case is None:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case를 찾을 수 없습니다."})
    facts = await repository.list_case_facts(case_id)
    questions = await repository.list_customer_questions(case_id)
    verifications = await repository.list_verifications(case_id)
    actions = await repository.list_actions(case_id)
    try:
        payload = await service.ai_client.build_case_support_snapshot(
            _case_support_ai_input(case_id, case, facts, questions, verifications, actions)
        )
        return to_public_case_support_snapshot(case_id, payload, available=True)
    except AiServiceError as exc:
        return PublicCaseSupportSnapshotResponse(case_id=case_id, available=False, warnings=[str(exc)])


@app.get("/api/cases/{case_id}/ai/case-support", response_model=PublicCaseSupportSnapshotResponse)
async def read_case_support_snapshot(case_id: str) -> PublicCaseSupportSnapshotResponse:
    return await get_case_support_snapshot(case_id)


@app.get("/api/cases/{case_id}/customer-question-candidates", response_model=list[PublicQuestionCandidateResponse])
async def list_customer_question_candidates(case_id: str) -> list[PublicQuestionCandidateResponse]:
    case = await repository.get(case_id)
    if case is None:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case를 찾을 수 없습니다."})
    queued = await repository.list_customer_questions(case_id)
    snapshot = await get_case_support_snapshot(case_id)
    # AI 장애 시에도 기존 deterministic 후보로 고객 확인 흐름을 멈추지 않는다.
    candidates = snapshot.recommended_questions if snapshot.available else build_customer_question_candidates(case, queued)
    # AI 결과를 그대로 신뢰하지 않는다. 질문 이력 기반 최종 중복 방지는 General API가 담당한다.
    return exclude_handled_question_candidates(candidates, queued)


@app.get("/api/cases/{case_id}/customer-questions", response_model=list[PublicCustomerQuestionResponse | PublicCustomerQuestionView])
async def list_customer_questions(case_id: str, view: Literal["bank", "customer"] = "bank") -> list[PublicCustomerQuestionResponse | PublicCustomerQuestionView]:
    await require_case(case_id)
    items = await repository.list_customer_questions(case_id)
    if view == "customer":
        return [to_public_customer_question_view(item) for item in items]
    return [to_public_customer_question(item) for item in items]


@app.post("/api/cases/{case_id}/customer-questions", response_model=list[PublicCustomerQuestionResponse], status_code=201)
async def queue_customer_questions(case_id: str, request: PublicQueueCustomerQuestionsRequest) -> list[PublicCustomerQuestionResponse]:
    await require_case(case_id)
    try:
        items = await repository.queue_customer_questions(
            case_id, [item.model_dump() for item in request.questions], request.requested_by
        )
    except KeyError as exc:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case를 찾을 수 없습니다."}) from exc
    await dispatch_next_customer_question_message(case_id)
    return [to_public_customer_question(item) for item in items]


async def _ensure_ai_customer_questions(
    case_id: str,
    snapshot: PublicCaseSupportSnapshotResponse | None = None,
) -> list[PublicCustomerQuestionResponse]:
    """Queue only allowlisted P0 safety questions and deliver one card at a time."""
    case = await repository.get(case_id)
    if case is None:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case를 찾을 수 없습니다."})
    existing = await repository.list_customer_questions(case_id)
    handled_fields = {normalize_target_field(item["target_field"]) for item in existing if item.get("status") in {"PENDING", "ASKED", "ANSWERED"}}
    facts = await repository.list_case_facts(case_id)
    handled_fields.update(
        normalize_target_field(str(item.get("field", "")))
        for item in facts
        if item.get("status") == "CONFIRMED"
    )
    if case.get("victim_transfer_status") in {"YES", "NO"}:
        handled_fields.add("transfer_status")
    if AUTONOMOUS_P0_QUESTION_FIELDS.issubset(handled_fields):
        await dispatch_next_customer_question_message(case_id)
        return []
    snapshot = snapshot or await get_case_support_snapshot(case_id)
    candidates = snapshot.recommended_questions if snapshot.available else build_customer_question_candidates(case, existing)
    p0_questions = [
        {**candidate.model_dump(mode="python"), "source": "CUSTOMER_AGENT"}
        for candidate in candidates
        if candidate.priority == "P0"
        and normalize_target_field(candidate.target_field) in AUTONOMOUS_P0_QUESTION_FIELDS
        and normalize_target_field(candidate.target_field) not in handled_fields
    ]
    if not p0_questions:
        await dispatch_next_customer_question_message(case_id)
        return []
    created = await repository.queue_customer_questions(case_id, p0_questions, "customer-agent")
    await dispatch_next_customer_question_message(case_id)
    return [to_public_customer_question(item) for item in created]


@app.post("/api/cases/{case_id}/ai/customer-questions/ensure", response_model=list[PublicCustomerQuestionResponse])
async def ensure_ai_customer_questions(case_id: str) -> list[PublicCustomerQuestionResponse]:
    return await _ensure_ai_customer_questions(case_id)


async def sync_ai_checklist_items(case_id: str, snapshot: PublicCaseSupportSnapshotResponse) -> list[dict]:
    """Persist each AI-recommended check once so unfinished staff work accumulates."""
    existing = await repository.list_actions(case_id)
    known_fields = {
        normalize_target_field(str(item.get("action_type", "")).split(":")[-1])
        for item in existing
        if str(item.get("action_type", "")).startswith(AI_CHECKLIST_ACTION_PREFIX)
    }
    candidates = [
        (normalize_target_field(item.target_field), item.priority, item.description)
        for item in snapshot.unresolved_items[:12]
    ] if snapshot.available else []
    facts = await repository.list_case_facts(case_id)
    candidates.extend(
        (
            normalize_target_field(str(item.get("field", ""))),
            "P0" if normalize_target_field(str(item.get("field", ""))) in AUTONOMOUS_P0_QUESTION_FIELDS else "P1",
            f"{CHECKLIST_FIELD_LABELS.get(normalize_target_field(str(item.get('field', ''))), '추가 확인 사항')}에 대한 고객 답변 “{item.get('value', '')}”을 사실로 확정할지 검토하세요.",
        )
        for item in facts
        if item.get("status") == "PROPOSED"
    )
    created: list[dict] = []
    for field, priority, description in candidates:
        if not field or field in known_fields:
            continue
        created.append(await repository.create_action(case_id, {
            "action_type": f"{AI_CHECKLIST_ACTION_PREFIX}{priority}:{field}",
            "actor_type": "SYSTEM",
            "note": description,
        }))
        known_fields.add(field)
    return created


async def dispatch_next_customer_question_message(case_id: str) -> PublicMessageResponse | None:
    """A confirmed queue is delivered one at a time through the public chat."""
    question = await repository.dispatch_next_customer_question(case_id)
    if question is None:
        return None
    message = await repository.append_message(case_id, {
        "actor_type": "CUSTOMER_AGENT", "actor_user_id": "customer-agent",
        "actor_display_name": "안전 상담 AI", "actor_role": "CUSTOMER_AGENT",
        "content": question["question_text"], "channel": "CUSTOMER", "audience": "CUSTOMER",
        "visibility": "CUSTOMER", "message_kind": "CHAT", "mentions": [], "log_event": False,
    })
    await repository.link_customer_question_message(case_id, question["question_id"], message["message_id"])
    question["question_message_id"] = message["message_id"]
    return to_public_message(message)


@app.post("/api/cases/{case_id}/customer-questions/{question_id}/answer", response_model=PublicCustomerQuestionResponse)
async def answer_customer_question(case_id: str, question_id: str, request: PublicAnswerCustomerQuestionRequest) -> PublicCustomerQuestionResponse:
    await require_case(case_id)
    try:
        message = await repository.append_message(case_id, {
            "actor_type": "CUSTOMER", "actor_user_id": request.actor_user_id,
            "actor_display_name": request.actor_display_name, "actor_role": "CUSTOMER",
            "content": request.raw_answer, "channel": "CUSTOMER", "audience": "CUSTOMER",
            "visibility": "CUSTOMER", "message_kind": "CHAT", "mentions": [], "log_event": False,
        })
        answered = await repository.answer_customer_question(case_id, question_id, message["message_id"], request.raw_answer)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail={"code": "CUSTOMER_QUESTION_NOT_FOUND", "message": "응답 대기 중인 질문을 찾을 수 없습니다."}) from exc
    await repository.propose_case_fact(case_id, question_id, request.raw_answer, message["message_id"])
    await repository.append_message(case_id, {
        "actor_type": "BANK_AGENT", "actor_user_id": "case-copilot",
        "actor_display_name": "CaseCopilot", "actor_role": "BANK_AGENT",
        "content": (
            "고객 답변 접수\n"
            f"질문: {answered.get('question_text', '고객 확인 질문')}\n"
            f"답변: {request.raw_answer}\n"
            "상태: 담당자 확인 전 정보 후보"
        ),
        "channel": "AI_INTERNAL", "audience": "BANK_INTERNAL", "visibility": "AI_PRIVATE",
        "message_kind": "SYSTEM_EVENT", "mentions": [], "private_owner_user_id": None,
        "reply_to_message_id": message["message_id"], "log_event": False,
    })
    if _answer_reports_customer_loss(answered, request.raw_answer):
        await _activate_customer_recovery(
            case_id,
            request.actor_user_id,
            request.actor_display_name,
            add_customer_acknowledgement=False,
        )
    await dispatch_next_customer_question_message(case_id)
    return to_public_customer_question(answered)


@app.get("/api/cases/{case_id}/facts", response_model=list[PublicCaseFactResponse])
async def list_case_facts(case_id: str) -> list[PublicCaseFactResponse]:
    await require_case(case_id)
    return [PublicCaseFactResponse.model_validate(item) for item in await repository.list_case_facts(case_id)]


@app.post("/api/cases/{case_id}/facts/{fact_id}/confirm", response_model=PublicCaseFactResponse)
async def confirm_case_fact(case_id: str, fact_id: str, request: PublicConfirmCaseFactRequest) -> PublicCaseFactResponse:
    await require_case(case_id)
    try:
        return PublicCaseFactResponse.model_validate(await repository.confirm_case_fact(case_id, fact_id, request.confirmed_by))
    except KeyError as exc:
        raise HTTPException(status_code=404, detail={"code": "CASE_FACT_NOT_FOUND", "message": "확정할 CaseFact를 찾을 수 없습니다."}) from exc


@app.get("/api/cases/{case_id}/personal-notes", response_model=list[PublicPersonalNoteResponse])
async def list_personal_notes(case_id: str, author_id: str) -> list[PublicPersonalNoteResponse]:
    await require_case(case_id)
    return [PublicPersonalNoteResponse.model_validate(item) for item in await repository.list_personal_notes(case_id, author_id)]


@app.post("/api/cases/{case_id}/personal-notes", response_model=PublicPersonalNoteResponse, status_code=201)
async def create_personal_note(case_id: str, request: PublicPersonalNoteCreateRequest) -> PublicPersonalNoteResponse:
    await require_case(case_id)
    return PublicPersonalNoteResponse.model_validate(await repository.create_personal_note(case_id, request.author_id, request.content))


@app.patch("/api/cases/{case_id}/personal-notes/{note_id}", response_model=PublicPersonalNoteResponse)
async def update_personal_note(case_id: str, note_id: str, request: PublicPersonalNoteUpdateRequest) -> PublicPersonalNoteResponse:
    await require_case(case_id)
    try:
        return PublicPersonalNoteResponse.model_validate(await repository.update_personal_note(case_id, note_id, request.author_id, request.content))
    except KeyError as exc:
        raise HTTPException(status_code=404, detail={"code": "PERSONAL_NOTE_NOT_FOUND", "message": "개인 메모를 찾을 수 없습니다."}) from exc


@app.delete("/api/cases/{case_id}/personal-notes/{note_id}", status_code=204)
async def delete_personal_note(case_id: str, note_id: str, author_id: str) -> None:
    await require_case(case_id)
    try:
        await repository.delete_personal_note(case_id, note_id, author_id)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail={"code": "PERSONAL_NOTE_NOT_FOUND", "message": "개인 메모를 찾을 수 없습니다."}) from exc


@app.post("/api/cases/{case_id}/attachments", response_model=PublicAttachmentResponse, status_code=201)
async def upload_case_attachment(
    case_id: str,
    request: Request,
    file_name: str,
    uploaded_by: str,
    visibility: Literal["BANK_INTERNAL", "CUSTOMER", "AI_PRIVATE"] = "CUSTOMER",
) -> PublicAttachmentResponse:
    await require_case(case_id)
    if not uploaded_by.strip():
        raise HTTPException(status_code=422, detail={"code": "UPLOADER_REQUIRED", "message": "업로드 사용자가 필요합니다."})
    content_length = request.headers.get("content-length")
    if content_length:
        try:
            if int(content_length) > MAX_ATTACHMENT_BYTES:
                raise HTTPException(status_code=413, detail={"code": "ATTACHMENT_TOO_LARGE", "message": f"파일은 최대 {MAX_ATTACHMENT_BYTES // (1024 * 1024)}MB까지 업로드할 수 있습니다."})
        except ValueError as exc:
            raise HTTPException(status_code=400, detail={"code": "INVALID_CONTENT_LENGTH", "message": "잘못된 파일 크기 정보입니다."}) from exc
    content = await request.body()
    if not content:
        raise HTTPException(status_code=422, detail={"code": "ATTACHMENT_EMPTY", "message": "빈 파일은 업로드할 수 없습니다."})
    if len(content) > MAX_ATTACHMENT_BYTES:
        raise HTTPException(status_code=413, detail={"code": "ATTACHMENT_TOO_LARGE", "message": f"파일은 최대 {MAX_ATTACHMENT_BYTES // (1024 * 1024)}MB까지 업로드할 수 있습니다."})

    original_name = _clean_file_name(file_name)
    mime_type, extension = _resolve_attachment_type(original_name, request.headers.get("content-type"))
    if not _has_valid_signature(mime_type, content):
        raise HTTPException(status_code=415, detail={"code": "ATTACHMENT_SIGNATURE_MISMATCH", "message": "파일 내용과 형식이 일치하지 않습니다."})

    case_directory = ATTACHMENT_STORAGE_ROOT / hashlib.sha256(case_id.encode("utf-8")).hexdigest()[:24]
    case_directory.mkdir(parents=True, exist_ok=True)
    stored_name = f"{uuid4().hex}{extension}"
    stored_path = (case_directory / stored_name).resolve()
    if not stored_path.is_relative_to(ATTACHMENT_STORAGE_ROOT):
        raise HTTPException(status_code=400, detail={"code": "INVALID_ATTACHMENT_PATH", "message": "잘못된 파일 경로입니다."})
    stored_path.write_bytes(content)
    try:
        record = await repository.create_attachment(case_id, {
            "original_name": original_name,
            "stored_name": stored_name,
            "storage_path": stored_path.relative_to(ATTACHMENT_STORAGE_ROOT).as_posix(),
            "mime_type": mime_type,
            "size_bytes": len(content),
            "sha256": hashlib.sha256(content).hexdigest(),
            "uploaded_by": uploaded_by.strip()[:80],
            "status": "UPLOADED",
            "visibility": visibility,
            "ai_readable": True,
            "created_at": datetime.now(timezone.utc).isoformat(),
        })
    except Exception:
        stored_path.unlink(missing_ok=True)
        raise
    return to_public_attachment(record)


@app.get("/api/cases/{case_id}/attachments", response_model=list[PublicAttachmentResponse])
async def list_case_attachments(case_id: str, view: Literal["bank", "customer"] = "bank") -> list[PublicAttachmentResponse]:
    await require_case(case_id)
    records = await repository.list_attachments(case_id)
    if view == "customer":
        records = [item for item in records if item.get("visibility") == "CUSTOMER"]
    return [to_public_attachment(item) for item in records]


@app.get("/api/internal/cases/{case_id}/attachments", response_model=list[PublicAttachmentResponse])
async def list_case_attachments_for_ai(case_id: str) -> list[PublicAttachmentResponse]:
    """AI 서비스용 메타데이터 목록. 서비스 인증은 배포 게이트웨이에서 적용한다."""
    await require_case(case_id)
    return [to_public_attachment(item, download_view="bank") for item in await repository.list_attachments(case_id) if item.get("ai_readable", True)]


@app.get("/api/cases/{case_id}/attachments/{attachment_id}/content")
async def download_case_attachment(case_id: str, attachment_id: str, view: Literal["bank", "customer"] = "customer") -> FileResponse:
    await require_case(case_id)
    record = await repository.get_attachment(case_id, attachment_id)
    if record is None:
        raise HTTPException(status_code=404, detail={"code": "ATTACHMENT_NOT_FOUND", "message": "첨부 파일을 찾을 수 없습니다."})
    if view == "customer" and record.get("visibility") != "CUSTOMER":
        raise HTTPException(status_code=403, detail={"code": "ATTACHMENT_FORBIDDEN", "message": "이 첨부 파일을 열 권한이 없습니다."})
    stored_path = (ATTACHMENT_STORAGE_ROOT / record["storage_path"]).resolve()
    if not stored_path.is_relative_to(ATTACHMENT_STORAGE_ROOT) or not stored_path.is_file():
        raise HTTPException(status_code=410, detail={"code": "ATTACHMENT_CONTENT_MISSING", "message": "첨부 파일 원본을 찾을 수 없습니다."})
    return FileResponse(stored_path, media_type=record["mime_type"], filename=None if record["mime_type"].startswith("image/") else record["original_name"])


@app.get("/api/cases/{case_id}/messages", response_model=list[PublicMessageResponse])
async def list_case_messages(case_id: str, channel: MessageChannel | None = None, view: Literal["bank", "customer"] = "bank") -> list[PublicMessageResponse]:
    await require_case(case_id)
    visible_channel = "CUSTOMER" if view == "customer" else channel
    messages = await repository.list_messages(case_id, visible_channel)
    if view == "customer":
        messages = [record for record in messages if record.get("visibility", record.get("audience")) == "CUSTOMER"]
    return [to_public_message(record) for record in messages]


def _customer_reports_loss(text: str) -> bool:
    """Route explicit customer loss statements to the recovery workflow.

    Negative expressions take precedence so phrases such as ``아직 송금 안 했어요``
    never activate recovery. Ambiguous messages remain in the normal AI chat.
    """
    compact = re.sub(r"\s+", "", text).lower()
    if not compact:
        return False
    if any(token in compact for token in (
        "송금안", "이체안", "보내지않", "송금하지않", "이체하지않", "아직안",
        "제공하지않", "알려주지않", "설치하지않", "아니요", "없어요", "없음",
    )):
        return False
    return any(token in compact for token in (
        "이미송금", "송금했", "이체했", "입금했", "돈을보냈", "돈보냈",
        "개인정보를제공", "개인정보알려", "계좌번호를알려", "주민번호를알려",
        "비밀번호를알려", "인증번호를알려", "otp를알려", "원격앱을설치", "원격제어앱설치",
        "사기당했", "피해를입었",
    ))


def _answer_reports_customer_loss(question: dict, raw_answer: str) -> bool:
    if _customer_reports_loss(raw_answer):
        return True
    target = normalize_target_field(str(question.get("target_field", "")))
    loss_fields = {
        "transfer_status", "personal_information_exposure",
        "authentication_information_exposure", "remote_control_app",
    }
    compact = re.sub(r"\s+", "", raw_answer).lower()
    return target in loss_fields and compact in {
        "예", "네", "있음", "있어요", "제공함", "제공했어요", "설치함", "설치했어요",
        "이미송금했어요", "이미이체했어요",
    }


async def _activate_customer_recovery(
    case_id: str,
    actor_user_id: str,
    actor_display_name: str,
    *,
    add_customer_acknowledgement: bool,
) -> PublicMessageResponse:
    """Idempotently activate recovery and publish one private bank alert."""
    case = await repository.get(case_id)
    if case is None:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case를 찾을 수 없습니다."})
    all_messages = await repository.list_messages(case_id)
    acknowledgement = next((item for item in all_messages
                            if item.get("channel") == "CUSTOMER"
                            and item.get("actor_user_id") == actor_user_id
                            and item.get("content") == "이미 사기 피해를 입었습니다. 피해구제 안내를 확인합니다."), None)
    if add_customer_acknowledgement and acknowledgement is None:
        await repository.append_message(case_id, {
            "actor_type": "CUSTOMER", "actor_user_id": actor_user_id,
            "actor_display_name": actor_display_name, "actor_role": "CUSTOMER",
            "content": "이미 사기 피해를 입었습니다. 피해구제 안내를 확인합니다.",
            "channel": "CUSTOMER", "audience": "CUSTOMER", "visibility": "CUSTOMER",
            "message_kind": "CHAT", "mentions": [], "log_event": False, "customer_emergency_ack": True,
        })
    alert = next((item for item in all_messages
                  if item.get("channel") == "AI_INTERNAL"
                  and item.get("actor_user_id") == "case-copilot"
                  and item.get("actor_display_name") == "CaseCopilot 긴급 알림"), None)
    if alert is None:
        alert = await repository.append_message(case_id, {
            "actor_type": "BANK_AGENT", "actor_user_id": "case-copilot",
            "actor_display_name": "CaseCopilot 긴급 알림", "actor_role": "BANK_AGENT",
            "content": "고객이 직접 사기 피해 발생을 신고했습니다. 피해 금액과 송금 정보를 확인하고 즉시 보호 조치를 검토해 주세요.",
            "channel": "AI_INTERNAL", "audience": "BANK_INTERNAL", "visibility": "AI_PRIVATE",
            "message_kind": "SYSTEM_EVENT", "mentions": ["CaseCopilot"],
            "private_owner_user_id": None, "log_event": False, "customer_emergency_alert": True,
            "customer_reported_by_user_id": actor_user_id,
            "customer_reported_by_display_name": actor_display_name,
        })
    if case.get("mode") != "RECOVERY" or case.get("victim_transfer_status") != "YES":
        latest = await repository.get(case_id) or case
        try:
            await repository.update_case(case_id, int(latest.get("version", 1)), {
                "victim_transfer_status": "YES", "mode": "RECOVERY",
            })
        except CaseVersionConflictError as exc:
            raise HTTPException(status_code=409, detail={"code": "VERSION_CONFLICT", "message": "Case가 변경되었습니다. 다시 시도해 주세요.", "current_version": exc.current_version}) from exc
    return to_public_message(alert)


@app.post("/api/cases/{case_id}/customer-emergency", response_model=PublicMessageResponse, status_code=201)
async def start_customer_emergency(case_id: str, request: PublicCustomerEmergencyRequest) -> PublicMessageResponse:
    """Persist one customer acknowledgement, one case-wide AI alert, and a linked Recovery event."""
    return await _activate_customer_recovery(
        case_id,
        request.actor_user_id,
        request.actor_display_name,
        add_customer_acknowledgement=True,
    )


@app.post("/api/cases/{case_id}/messages", response_model=PublicMessageResponse, status_code=201)
async def create_case_message(case_id: str, request: PublicCreateMessageRequest) -> PublicMessageResponse:
    await require_case(case_id)
    if request.client_request_id:
        existing = await repository.find_message_by_client_request_id(case_id, request.client_request_id)
        if isinstance(existing, dict):
            return to_public_message(existing)
    try:
        record = await repository.append_message(case_id, request.model_dump())
    except KeyError as exc:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case를 찾을 수 없습니다."}) from exc
    except ValueError as exc:
        raise HTTPException(status_code=422, detail={"code": "ATTACHMENT_NOT_FOUND", "message": "메시지에 연결할 첨부 파일을 찾을 수 없습니다."}) from exc
    if (
        request.actor_type == "CUSTOMER"
        and request.channel == "CUSTOMER"
        and _customer_reports_loss(request.content)
    ):
        await _activate_customer_recovery(
            case_id,
            request.actor_user_id,
            request.actor_display_name,
            add_customer_acknowledgement=False,
        )
    return to_public_message(record)


@app.get("/api/cases/{case_id}/events", response_model=list[PublicCaseEventResponse])
async def list_case_events(case_id: str, after: int | None = None) -> list[PublicCaseEventResponse]:
    await require_case(case_id)
    return [to_public_event(record) for record in await repository.list_events(case_id, after)]


@app.get("/api/cases/{case_id}/members", response_model=list[PublicCaseMemberResponse])
async def list_case_members(case_id: str) -> list[PublicCaseMemberResponse]:
    await require_case(case_id)
    return [PublicCaseMemberResponse.model_validate(item) for item in await repository.list_members(case_id)]


@app.post("/api/cases/{case_id}/members", response_model=PublicCaseMemberResponse, status_code=201)
async def upsert_case_member(case_id: str, request: PublicCaseMemberUpsertRequest) -> PublicCaseMemberResponse:
    await require_case(case_id)
    try:
        return PublicCaseMemberResponse.model_validate(await repository.upsert_member(case_id, request.model_dump()))
    except KeyError as exc:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case를 찾을 수 없습니다."}) from exc


@app.put("/api/cases/{case_id}/assignee", response_model=PublicPrimaryAssigneeResponse)
async def set_case_primary_assignee(case_id: str, request: PublicPrimaryAssigneeRequest) -> PublicPrimaryAssigneeResponse:
    try:
        display_name = await repository.set_primary_assignee(case_id, request.display_name)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case not found."}) from exc
    return PublicPrimaryAssigneeResponse(case_id=case_id, display_name=display_name)


@app.get("/api/cases/{case_id}/presence", response_model=list[PublicCasePresenceResponse])
async def list_case_presence(case_id: str) -> list[PublicCasePresenceResponse]:
    await require_case(case_id)
    return [PublicCasePresenceResponse.model_validate(item) for item in await repository.list_presence(case_id)]


@app.post("/api/cases/{case_id}/presence/heartbeat", response_model=PublicCasePresenceResponse)
async def heartbeat_case_presence(case_id: str, request: PublicPresenceHeartbeatRequest) -> PublicCasePresenceResponse:
    await require_case(case_id)
    try:
        return PublicCasePresenceResponse.model_validate(await repository.heartbeat_presence(case_id, request.model_dump()))
    except KeyError as exc:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case를 찾을 수 없습니다."}) from exc


def build_mvp_copilot_reply(case: dict, verifications: list[dict], prompt: str) -> str:
    unresolved = [item.get("claim", "추가 확인 항목") for item in verifications if item.get("status") != "COMPLETED"]
    unresolved_text = ", ".join(unresolved[:3]) or "추가 확인 항목이 아직 등록되지 않았습니다."
    return (
        f"[MVP CaseCopilot] 요청: {prompt.strip()}\n\n"
        f"현재 Case는 {case.get('status', 'TRIAGE')} 상태이며, 요약은 “{case.get('initial_brief', '확인안됨')}”입니다.\n"
        f"미완료 검증: {unresolved_text}\n"
        "다음 단계로 고객에게 확인할 사실을 한 번에 하나씩 정리하고, 고객 전송 전 담당자 승인을 받으세요."
    )


@app.post("/api/cases/{case_id}/ai/work-cards", response_model=CaseWorkCardOutput)
async def generate_case_work_card(case_id: str, request: PublicWorkCardGenerateRequest) -> CaseWorkCardOutput:
    case = await repository.get(case_id)
    if case is None:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case를 찾을 수 없습니다."})
    facts = await repository.list_case_facts(case_id)
    verifications = await repository.list_verifications(case_id)
    actions = await repository.list_actions(case_id)
    messages = await repository.list_messages(case_id)
    attachments = await repository.list_attachments(case_id)
    support = await get_case_support_snapshot(case_id)
    try:
        payload = await service.ai_client.generate_work_card({
            "case_id": case_id,
            "card_type": request.card_type,
            "case_summary": case.get("initial_brief", ""),
            "workflow_status": case.get("status", "TRIAGE"),
            "case_mode": case.get("mode", "PREVENT"),
            "fraud_type": case.get("fraud_type"),
            "known_facts": [f"{item.get('field')}: {item.get('value')} ({item.get('status')})" for item in facts[:30]],
            "recent_conversation": [
                f"{item.get('actor_display_name', item.get('actor_type', '작성자'))}: {item.get('content', '')[:500]}"
                for item in messages[-20:]
                if item.get("channel") in {"TEAM", "CUSTOMER"}
            ],
            "pending_actions": [
                f"{item.get('action_type')}: {item.get('note') or '상세 내용 없음'} ({item.get('status', 'REQUESTED')})"
                for item in actions_for_ai(actions) if item.get("status") not in {"COMPLETED", "CANCELLED"}
            ][:20],
            "attachment_summaries": [
                f"{item.get('original_name', '첨부 파일')} ({item.get('mime_type', '형식 미상')}, {item.get('visibility', '공개 범위 미상')})"
                for item in attachments[-10:]
            ],
            "unresolved_items": [f"{item.priority}: {item.description}" for item in support.unresolved_items[:20]],
            "pending_verifications": [f"{item.get('target')}: {item.get('claim')}" for item in verifications if item.get("status") != "COMPLETED"][:20],
            "question_candidates": [item.model_dump(mode="python") for item in support.recommended_questions[:10]],
        })
        return CaseWorkCardOutput.model_validate(payload)
    except AiServiceQuotaError as exc:
        raise HTTPException(status_code=429, detail={"code": "OPENAI_QUOTA_EXHAUSTED", "message": str(exc)}) from exc
    except AiServiceAuthenticationError as exc:
        raise HTTPException(status_code=401, detail={"code": "OPENAI_AUTHENTICATION_FAILED", "message": str(exc)}) from exc
    except AiServiceError as exc:
        raise HTTPException(status_code=503, detail={"code": "AI_WORK_CARD_FAILED", "message": str(exc)}) from exc


@app.post("/api/cases/{case_id}/ai/invocations", response_model=PublicAiInvocationResponse, status_code=201)
async def invoke_case_copilot(case_id: str, request: PublicAiInvocationRequest) -> PublicAiInvocationResponse:
    case = await repository.get(case_id)
    if case is None:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case를 찾을 수 없습니다."})
    verifications = await repository.list_verifications(case_id)
    facts = await repository.list_case_facts(case_id)
    actions = await repository.list_actions(case_id)
    attachments = await repository.list_attachments(case_id)
    all_messages = await repository.list_messages(case_id)
    members = await repository.list_members(case_id)
    primary_assignee = next(
        (item.get("display_name") for item in members if item.get("role") == "CASE_OWNER"),
        None,
    )
    member_role_labels = {
        "CASE_OWNER": "메인 담당자",
        "CHAT_OPERATOR": "상담 담당자",
        "REVIEWER": "검토자",
        "VIEWER": "열람자",
    }
    unresolved = [item.get("claim", "추가 확인 항목") for item in verifications if item.get("status") != "COMPLETED"]
    try:
        ai_reply = await service.ai_client.generate_case_copilot_reply({
            "case_id": case_id,
            "prompt": request.prompt,
            "case_summary": case.get("initial_brief", ""),
            "workflow_status": case.get("status", "TRIAGE"),
            "fraud_type": case.get("fraud_type"),
            "transfer_status": case.get("victim_transfer_status"),
            "primary_assignee": primary_assignee,
            "participants": [
                f"{item.get('display_name', '이름 미상')} ({member_role_labels.get(item.get('role'), item.get('role', '역할 미상'))})"
                for item in members
            ][:30],
            "known_facts": [f"{item.get('field')}: {item.get('value')} ({item.get('status')})" for item in facts[:30]],
            "recent_conversation": [
                f"{item.get('actor_display_name', item.get('actor_type', '작성자'))}: {item.get('content', '')[:500]}"
                for item in all_messages[-30:]
                if item.get("channel") in {"TEAM", "CUSTOMER"}
                or (item.get("channel") == "AI_INTERNAL" and item.get("private_owner_user_id") in {None, request.requester_user_id})
            ][-20:],
            "pending_actions": [
                f"{item.get('action_type')}: {item.get('note') or '상세 내용 없음'} ({item.get('status', 'REQUESTED')})"
                for item in actions_for_ai(actions) if item.get("status") not in {"COMPLETED", "CANCELLED"}
            ][:20],
            "attachment_summaries": [
                f"{item.get('original_name', '첨부 파일')} ({item.get('mime_type', '형식 미상')}, {item.get('visibility', '공개 범위 미상')})"
                for item in attachments[-10:]
            ],
            "unresolved_verifications": unresolved[:10],
            "assistant_mode": "BANK_INTERNAL",
            "customer_progress": progress_ai_context(build_customer_progress(actions)),
            "response_style": request.response_style,
        })
    except AiServiceQuotaError as exc:
        raise HTTPException(status_code=429, detail={"code": "OPENAI_QUOTA_EXHAUSTED", "message": str(exc)}) from exc
    except AiServiceAuthenticationError as exc:
        raise HTTPException(status_code=401, detail={"code": "OPENAI_AUTHENTICATION_FAILED", "message": str(exc)}) from exc
    except AiServiceError as exc:
        raise HTTPException(status_code=503, detail={"code": "AI_CASE_COPILOT_FAILED", "message": str(exc)}) from exc
    content = ai_reply["content"]
    is_team_request = request.channel == "TEAM"
    message = await repository.append_message(case_id, {
        "actor_type": "BANK_AGENT", "actor_user_id": "case-copilot", "actor_display_name": "CaseCopilot",
        "actor_role": "BANK_AGENT", "content": content,
        "channel": "TEAM" if is_team_request else "AI_INTERNAL", "audience": "BANK_INTERNAL",
        "visibility": "BANK_INTERNAL" if is_team_request else "AI_PRIVATE", "message_kind": "AI_RESPONSE", "mentions": ["CaseCopilot"],
        "private_owner_user_id": None if is_team_request else request.requester_user_id, "client_request_id": request.client_request_id, "log_event": True,
    })
    return PublicAiInvocationResponse(
        invocation_id=f"ai-{uuid4().hex}", message_id=message["message_id"], case_id=case_id,
        channel="TEAM" if is_team_request else "AI_INTERNAL", content=content, model_mode=ai_reply["model_mode"], created_at=message["created_at"],
    )


@app.post("/api/cases/{case_id}/ai/customer-replies", response_model=PublicMessageResponse, status_code=201)
async def invoke_customer_support_ai(case_id: str, request: PublicCustomerAiReplyRequest) -> PublicMessageResponse:
    """Generate a customer-safe reply without exposing the bank-only Case bundle."""
    case = await repository.get(case_id)
    if case is None:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case를 찾을 수 없습니다."})
    all_messages = await repository.list_messages(case_id, "CUSTOMER")
    questions = await repository.list_customer_questions(case_id)
    progress = build_customer_progress(await repository.list_actions(case_id))
    verifications = await repository.list_verifications(case_id)
    published_results = [f"{item.get('target')}: {item.get('result_summary')} (공개 확인 결과)"
                         for item in verifications if item.get('customer_visible') and item.get('status') == 'COMPLETED' and item.get('result_summary')][-10:]
    attachments = await repository.list_attachments(case_id)
    customer_history = [
        f"{item.get('actor_display_name', '상담 참여자')}: {item.get('content', '')[:500]}"
        for item in all_messages[-20:]
        if item.get("visibility") == "CUSTOMER"
    ]
    answered = [
        f"질문: {item.get('question_text', '')} / 고객 답변: {item.get('answer_text', '')}"
        for item in questions
        if item.get("status") == "ANSWERED" and item.get("answer_text")
    ][-10:]
    try:
        ai_reply = await service.ai_client.generate_case_copilot_reply({
            "case_id": case_id,
            "prompt": request.prompt,
            "case_summary": "",
            "workflow_status": case.get("status", "TRIAGE"),
            "fraud_type": case.get("fraud_type"),
            "transfer_status": case.get("victim_transfer_status"),
            "known_facts": answered,
            "recent_conversation": customer_history,
            "pending_actions": [],
            "unresolved_verifications": [],
            "assistant_mode": "CUSTOMER_SUPPORT",
            "primary_assignee": case.get('primary_assignee'),
            "customer_progress": progress_ai_context(progress),
            "published_verification_results": published_results,
            "attachment_summaries": [str(item.get('original_name', '첨부 자료')) for item in attachments
                                     if item.get('visibility') == 'CUSTOMER'][-10:],
        })
    except AiServiceQuotaError as exc:
        raise HTTPException(status_code=429, detail={"code": "OPENAI_QUOTA_EXHAUSTED", "message": str(exc)}) from exc
    except AiServiceAuthenticationError as exc:
        raise HTTPException(status_code=401, detail={"code": "OPENAI_AUTHENTICATION_FAILED", "message": str(exc)}) from exc
    except AiServiceError as exc:
        raise HTTPException(status_code=503, detail={"code": "AI_CUSTOMER_SUPPORT_FAILED", "message": str(exc)}) from exc
    message = await repository.append_message(case_id, {
        "actor_type": "CUSTOMER_AGENT", "actor_user_id": "customer-agent",
        "actor_display_name": "안전 상담 AI", "actor_role": "CUSTOMER_AGENT",
        "content": ai_reply["content"], "channel": "CUSTOMER", "audience": "CUSTOMER",
        "visibility": "CUSTOMER", "message_kind": "AI_RESPONSE", "mentions": [],
        "reply_to_message_id": request.reply_to_message_id, "client_request_id": request.client_request_id,
        "log_event": False,
    })
    return to_public_message(message)


async def run_proactive_case_automation(case_id: str) -> bool:
    """Reconcile safety-critical AI questions without a frontend request.

    The automation is deliberately fail-open for the rest of the Case API:
    an AI outage must not prevent staff or customer messages from being saved.
    """
    try:
        snapshot = await get_case_support_snapshot(case_id)
        await _ensure_ai_customer_questions(case_id, snapshot)
        await sync_ai_checklist_items(case_id, snapshot)
        return True
    except HTTPException as exc:
        if exc.status_code != 404:
            logger.warning("Proactive question reconciliation failed for %s: %s", case_id, exc.detail)
    except Exception as exc:
        error_number = exc.args[0] if exc.args and isinstance(exc.args[0], int) else None
        logger.warning("Proactive question reconciliation failed for %s: %s errno=%s", case_id, type(exc).__name__, error_number)
    return False


async def proactive_case_worker() -> None:
    """Observe durable Case revisions and reconcile automation only on changes."""
    while True:
        try:
            await reconcile_changed_cases_once()
        except asyncio.CancelledError:
            raise
        except Exception as exc:
            logger.warning("Proactive Case worker iteration failed: %s", type(exc).__name__)
        await asyncio.sleep(PROACTIVE_CASE_POLL_SECONDS)


async def reconcile_changed_cases_once() -> int:
    """Run one durable-revision scan; exposed separately for contract tests."""
    reconciled = 0
    for case in await repository.list():
        case_id = str(case.get("case_id", ""))
        if not case_id or case.get("status") == "CLOSED" or case.get("mode") == "CLOSED":
            continue
        revision = str(case.get("updated_at", ""))
        if revision and _proactive_case_revisions.get(case_id) == revision:
            continue
        if await run_proactive_case_automation(case_id):
            latest = await repository.get(case_id)
            _proactive_case_revisions[case_id] = str((latest or case).get("updated_at", revision))
            reconciled += 1
    return reconciled


@app.on_event("startup")
async def start_proactive_case_worker() -> None:
    global _proactive_worker_task
    ping = getattr(repository, "ping", None)
    if callable(ping):
        await ping()
    if os.getenv("PROACTIVE_QUESTION_AUTOMATION", "1").lower() not in {"0", "false", "off"}:
        _proactive_worker_task = asyncio.create_task(proactive_case_worker())


@app.on_event("shutdown")
async def stop_proactive_case_worker() -> None:
    global _proactive_worker_task
    if _proactive_worker_task is not None:
        _proactive_worker_task.cancel()
        try:
            await _proactive_worker_task
        except asyncio.CancelledError:
            pass
        _proactive_worker_task = None
    close = getattr(repository, "close", None)
    if callable(close):
        await close()


@app.post("/api/cases/{case_id}/ai/messages/{message_id}/share", response_model=PublicMessageResponse, status_code=201)
async def share_ai_message_to_team(case_id: str, message_id: str, request: PublicAiShareRequest) -> PublicMessageResponse:
    await require_case(case_id)
    source = next((item for item in await repository.list_messages(case_id, "AI_INTERNAL") if item.get("message_id") == message_id), None)
    if source is None or source.get("actor_type") != "BANK_AGENT" or source.get("message_kind") != "AI_RESPONSE":
        raise HTTPException(status_code=404, detail={"code": "AI_MESSAGE_NOT_FOUND", "message": "공유할 AI 답변을 찾을 수 없습니다."})
    shared = await repository.append_message(case_id, {
        "actor_type": "BANK_AGENT", "actor_user_id": source.get("actor_user_id", "case-copilot"),
        "actor_display_name": source.get("actor_display_name", "CaseCopilot"), "actor_role": "BANK_AGENT",
        "content": source["content"], "channel": "TEAM", "audience": "BANK_INTERNAL",
        "visibility": "BANK_INTERNAL", "message_kind": "AI_RESPONSE", "mentions": ["CaseCopilot"],
        "reply_to_message_id": message_id, "shared_by_user_id": request.shared_by_user_id,
        "shared_by_display_name": request.shared_by_display_name, "log_event": True,
    })
    return to_public_message(shared)


@app.patch("/api/cases/{case_id}/verifications/{verification_task_id}", response_model=PublicVerificationResponse)
async def update_case_verification(case_id: str, verification_task_id: str, request: PublicUpdateVerificationRequest) -> PublicVerificationResponse:
    await require_case(case_id)
    try:
        return to_public_verification(await repository.update_verification(case_id, verification_task_id, request.expected_version, request.status, request.model_dump(exclude={"expected_version", "status"}, exclude_none=True)))
    except CaseVersionConflictError as exc:
        raise HTTPException(status_code=409, detail={"code": "VERSION_CONFLICT", "message": "Verification task has changed.", "current_version": exc.current_version}) from exc
    except KeyError as exc:
        raise HTTPException(status_code=404, detail={"code": "VERIFICATION_NOT_FOUND", "message": "Verification task not found."}) from exc


@app.get("/api/cases/{case_id}/verifications", response_model=list[PublicVerificationResponse])
async def list_case_verifications(case_id: str) -> list[PublicVerificationResponse]:
    await require_case(case_id)
    return [to_public_verification(record) for record in await repository.list_verifications(case_id)]


@app.post("/api/cases/{case_id}/verifications", response_model=PublicVerificationResponse, status_code=201)
async def create_case_verification(case_id: str, request: PublicCreateVerificationRequest) -> PublicVerificationResponse:
    await require_case(case_id)
    try:
        return to_public_verification(await repository.create_verification(case_id, request.model_dump()))
    except KeyError as exc:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case를 찾을 수 없습니다."}) from exc


@app.get("/api/cases/{case_id}/actions", response_model=list[PublicActionResponse])
async def list_case_actions(case_id: str) -> list[PublicActionResponse]:
    await require_case(case_id)
    return [to_public_action(record) for record in await repository.list_actions(case_id) if not record['action_type'].startswith(PROGRESS_PREFIX)]


@app.get('/api/cases/{case_id}/customer-progress', response_model=list[CustomerProgressItem])
async def get_customer_progress(case_id: str):
    await require_case(case_id)
    return build_customer_progress(await repository.list_actions(case_id))


class DisplayEditRequest(BaseModel):
    expected_version: int = Field(ge=0)
    operation: Literal['EDIT', 'DELETE', 'RESTORE', 'RESET']
    text: str | None = Field(default=None, max_length=4000)

    @model_validator(mode='after')
    def validate_change(self):
        ContextItemChange(expected_version=max(1, self.expected_version), operation=self.operation, text=self.text)
        return self


async def context_display_repository(case_id):
    await require_case(case_id)
    return ContextItemRepository(repository) if isinstance(repository, MySqlCaseRepository) else InMemoryContextItemRepository(repository)


async def bank_display_repository(case_id, actor_user_id):
    store = await context_display_repository(case_id)
    members = await repository.list_members(case_id)
    if not any(item.get('user_id') == actor_user_id and item.get('role') in {'CASE_OWNER', 'CHAT_OPERATOR', 'REVIEWER'} for item in members):
        raise HTTPException(status_code=403, detail='이 사건의 담당자 또는 검토자만 맥락을 편집할 수 있습니다.')
    return store


@app.get('/api/cases/{case_id}/context-display')
async def read_context_display(case_id: str, actor_user_id: str):
    # 사건 화면을 여는 순간 프론트가 현재 직원을 참여자로 등록한다. 새 Case에서는
    # 그 등록과 이 조회가 동시에 도착할 수 있다. 아직 역할이 준비되지 않았다면
    # 오류나 다른 직원의 수정본을 내보내지 않고 빈 편집본으로 응답한다.
    # 실제 변경(PATCH)은 아래에서 계속 역할을 검증한다.
    store = await context_display_repository(case_id)
    members = await repository.list_members(case_id)
    if not any(item.get('user_id') == actor_user_id and item.get('role') in {'CASE_OWNER', 'CHAT_OPERATOR', 'REVIEWER'} for item in members):
        return []
    return [item for item in await store.list_items(case_id, include_deleted=True) if item.semantic_key == 'display']


@app.patch('/api/cases/{case_id}/context-display/{section}')
async def edit_context_display(case_id: str, section: Section, actor_user_id: str, request: DisplayEditRequest):
    store = await bank_display_repository(case_id, actor_user_id)
    try:
        return await store.edit_section(case_id, section, request.expected_version, request.operation, request.text, actor_user_id)
    except ContextItemConflictError as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc


@app.put('/api/cases/{case_id}/customer-progress/{step}', response_model=list[CustomerProgressItem])
async def update_customer_progress(case_id: str, step: ProgressStep, request: UpdateCustomerProgress):
    await require_case(case_id)
    try:
        await repository.create_action(case_id, {'_progress_command': {
            'step': step, 'request_confirmation': False, 'values': request.model_dump(mode='json'),
        }})
    except ProgressConflict as exc:
        raise HTTPException(status_code=409, detail={'code': 'PROGRESS_CONFLICT', 'message': str(exc)}) from exc
    return build_customer_progress(await repository.list_actions(case_id))


@app.post('/api/cases/{case_id}/customer-progress/{step}/confirmation-request', response_model=list[CustomerProgressItem])
async def request_progress_confirmation(case_id: str, step: ProgressStep):
    await require_case(case_id)
    await repository.create_action(case_id, {'_progress_command': {'step': step, 'request_confirmation': True}})
    return build_customer_progress(await repository.list_actions(case_id))


@app.post("/api/cases/{case_id}/actions", response_model=PublicActionResponse, status_code=201)
async def create_case_action(case_id: str, request: PublicCreateActionRequest) -> PublicActionResponse:
    await require_case(case_id)
    if request.action_type.startswith(PROGRESS_PREFIX):
        raise HTTPException(status_code=422, detail='고객 처리 상태는 전용 처리 결과 화면에서 기록해 주세요.')
    try:
        return to_public_action(await repository.create_action(case_id, request.model_dump()))
    except KeyError as exc:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case를 찾을 수 없습니다."}) from exc


@app.patch("/api/cases/{case_id}/actions/{action_id}", response_model=PublicActionResponse)
async def update_case_action(case_id: str, action_id: str, request: PublicUpdateActionRequest) -> PublicActionResponse:
    await require_case(case_id)
    actions = await repository.list_actions(case_id)
    current = next((item for item in actions if item['action_id'] == action_id), None)
    if current and current['action_type'].startswith(PROGRESS_PREFIX):
        raise HTTPException(status_code=422, detail='고객 처리 상태 이력은 일반 체크리스트로 변경할 수 없습니다.')
    try:
        if current is None and request.status is None:
            raise KeyError(action_id)
        next_status = request.status or current.get('status', 'REQUESTED')
        if request.note is None:
            updated = await repository.update_action(case_id, action_id, next_status, request.updated_by)
        else:
            updated = await repository.update_action(case_id, action_id, next_status, request.updated_by, request.note)
        return to_public_action(updated)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail={"code": "CASE_ACTION_NOT_FOUND", "message": "체크리스트 항목을 찾을 수 없습니다."}) from exc


async def create_case_control_action(case_id: str, action_type: str, request: PublicActionCommandRequest) -> PublicActionResponse:
    await require_case(case_id)
    try:
        return to_public_action(await repository.create_action(case_id, {"action_type": action_type, "actor_type": "BANK_STAFF", "note": request.note}))
    except KeyError as exc:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case not found."}) from exc


@app.post("/api/cases/{case_id}/takeover", response_model=PublicActionResponse, status_code=201)
async def start_human_takeover(case_id: str, request: PublicActionCommandRequest) -> PublicActionResponse:
    return await create_case_control_action(case_id, "HUMAN_TAKEOVER", request)


@app.post("/api/cases/{case_id}/resume", response_model=PublicActionResponse, status_code=201)
async def resume_ai(case_id: str, request: PublicActionCommandRequest) -> PublicActionResponse:
    return await create_case_control_action(case_id, "RESUME_AI", request)


@app.get("/api/cases/{case_id}/bundle", response_model=PublicCaseBundleResponse)
async def get_case_bundle(case_id: str, view: Literal["entry", "customer", "bank"] = "entry") -> PublicCaseBundleResponse:
    record = await repository.get(case_id)
    if record is None:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case를 찾을 수 없습니다."})
    # A customer bundle must never carry staff or AI-internal messages. The
    # dedicated messages endpoint applies the same channel constraint.
    visible_channel = "CUSTOMER" if view == "customer" else None
    messages = [to_public_message(item).model_dump(mode="json") for item in await repository.list_messages(case_id, visible_channel)]
    action_records = await repository.list_actions(case_id)
    customer_progress = build_customer_progress(action_records)
    actions = [to_public_action(item) for item in action_records if not item['action_type'].startswith(PROGRESS_PREFIX)]
    verifications = [to_public_verification(item) for item in await repository.list_verifications(case_id)]
    question_records = await repository.list_customer_questions(case_id)
    questions = [to_public_customer_question(item).model_dump(mode="json") for item in question_records]
    progress_items = [
        {"key": "customer_questions", "total": len(question_records), "answered": sum(item.get("status") == "ANSWERED" for item in question_records)},
        {"key": "verification_tasks", "total": len(verifications), "completed": sum(item.status == "COMPLETED" for item in verifications)},
    ]
    customer_verification_results: list[PublicCustomerVerificationResult] = []
    events = [to_public_event(item).model_dump(mode="json") for item in await repository.list_events(case_id)]
    # Customer view hides event details, but still needs the Case activity cursor
    # to identify the latest synchronized state.
    cursor = str(events[-1]["event_id"]) if events else None
    voice = await repository.get_voice_session(case_id)
    if view == "customer":
        questions = [to_public_customer_question_view(item).model_dump(mode="json") for item in question_records]
        customer_verification_results = [
            PublicCustomerVerificationResult(
                verification_task_id=item.verification_task_id,
                target=item.target,
                result_summary=item.result_summary,
                published_at=item.updated_at,
            )
            for item in verifications
            if item.status == "COMPLETED" and item.customer_visible and item.result_summary
        ]
        messages = [item for item in messages if item.get("visibility") == "CUSTOMER"]
        actions, verifications, events, voice = [], [], [], None
    return PublicCaseBundleResponse(
        customer_progress=customer_progress,
        case=to_public_case_summary_response(record).model_dump(mode="json"),
        live_report=None if view == "customer" else record.get("initial_report"),
        questions=questions, progress_items=progress_items, verification_tasks=verifications,
        customer_verification_results=customer_verification_results,
        recent_messages=messages[-50:], recent_actions=actions if view == "bank" else actions[-50:], recent_events=events[-50:],
        voice_session=PublicVoiceSessionResponse.model_validate(voice) if voice else None, cursor=cursor,
    )


@app.post("/api/cases/{case_id}/voice-sessions", response_model=PublicVoiceSessionResponse, status_code=201)
async def create_voice_session(case_id: str, request: PublicCreateVoiceSessionRequest) -> PublicVoiceSessionResponse:
    await require_case(case_id)
    try:
        return PublicVoiceSessionResponse.model_validate(await repository.create_voice_session(case_id, request.participants))
    except KeyError as exc:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case not found."}) from exc


@app.patch("/api/cases/{case_id}/voice-sessions/{session_id}", response_model=PublicVoiceSessionResponse)
async def update_voice_session(case_id: str, session_id: str, request: PublicUpdateVoiceSessionRequest) -> PublicVoiceSessionResponse:
    await require_case(case_id)
    try:
        return PublicVoiceSessionResponse.model_validate(await repository.update_voice_session(case_id, session_id, request.status))
    except KeyError as exc:
        raise HTTPException(status_code=404, detail={"code": "VOICE_SESSION_NOT_FOUND", "message": "Voice session not found."}) from exc


@app.get("/api/cases/{case_id}/voice-sessions/{session_id}/transcript", response_model=list[PublicTranscriptResponse])
async def list_voice_transcript(case_id: str, session_id: str) -> list[PublicTranscriptResponse]:
    await require_case(case_id)
    return [PublicTranscriptResponse.model_validate(item) for item in await repository.list_transcript(case_id, session_id)]


@app.post("/api/cases/{case_id}/voice-sessions/{session_id}/transcript", response_model=PublicTranscriptResponse, status_code=201)
async def append_voice_transcript(case_id: str, session_id: str, request: PublicCreateTranscriptRequest) -> PublicTranscriptResponse:
    await require_case(case_id)
    try:
        return PublicTranscriptResponse.model_validate(await repository.append_transcript(case_id, session_id, request.model_dump()))
    except KeyError as exc:
        raise HTTPException(status_code=404, detail={"code": "VOICE_SESSION_NOT_FOUND", "message": "Voice session not found."}) from exc


@app.get("/api/cases/{case_id}/reports/live")
async def get_live_report(case_id: str) -> dict:
    record = await repository.get(case_id)
    if record is None:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case를 찾을 수 없습니다."})
    return record["initial_report"]


@app.post("/api/cases/{case_id}/reports/finalize", response_model=PublicReportResponse)
async def finalize_case_report(case_id: str, request: AdminCaseFinalizeRequest) -> PublicReportResponse:
    require_admin_password(request.password)
    try:
        case = await repository.get(case_id)
        if case is None:
            raise KeyError(case_id)
        facts, verifications, actions, messages, questions = await asyncio.gather(
            repository.list_case_facts(case_id),
            repository.list_verifications(case_id),
            repository.list_actions(case_id),
            repository.list_messages(case_id),
            repository.list_customer_questions(case_id),
        )
        ai_report = await service.ai_client.generate_final_report({
            "case_id": case_id,
            "case_summary": case.get("initial_brief", ""),
            "workflow_status": case.get("status", "TRIAGE"),
            "case_mode": case.get("mode", "PREVENT"),
            "known_facts": [f"{item.get('field')}: {item.get('value')} ({item.get('status')})" for item in facts[:40]],
            "recent_conversation": [
                f"{item.get('actor_display_name', item.get('actor_type', '작성자'))}: {item.get('content', '')[:500]}"
                for item in messages[-30:] if item.get("message_kind") != "REPORT_CARD"
            ],
            "verification_results": [
                f"{item.get('target')}: {item.get('result_summary') or item.get('claim')} ({item.get('status')})"
                for item in verifications[:30]
            ],
            "action_results": [
                f"{item.get('action_type')}: {item.get('note') or '상세 내용 없음'} ({item.get('status')})"
                for item in actions_for_ai(actions)[:30]
            ],
            "customer_answers": [
                f"질문: {item.get('question_text')} / 고객 답변: {item.get('answer_text')}"
                for item in questions if item.get("answer_text")
            ][:30],
            "closure_note": request.note,
        })
        report_card = {
            "title": ai_report["title"],
            "executive_summary": ai_report["executive_summary"],
            "incident_summary": ai_report["incident_summary"],
            "verified_facts": ai_report.get("verified_facts", []),
            "actions_taken": ai_report.get("actions_taken", []),
            "resolution": ai_report["resolution"],
            "follow_up": ai_report.get("follow_up", []),
            "cautions": ai_report.get("cautions", []),
            "model_mode": ai_report.get("model_mode", "AI"),
        }
        sections = [
            {"section_key": "executive_summary", "content": {"text": report_card["executive_summary"]}},
            {"section_key": "incident_summary", "content": {"text": report_card["incident_summary"]}},
            {"section_key": "verified_facts", "content": {"items": report_card["verified_facts"]}},
            {"section_key": "actions_taken", "content": {"items": report_card["actions_taken"]}},
            {"section_key": "resolution", "content": {"text": report_card["resolution"], "closure_note": request.note}},
            {"section_key": "follow_up", "content": {"items": report_card["follow_up"]}},
            {"section_key": "cautions", "content": {"items": report_card["cautions"]}},
        ]
        report = await repository.finalize_report(case_id, request.expected_version, request.note, sections, report_card)
    except AiServiceQuotaError as exc:
        raise HTTPException(status_code=429, detail={"code": "OPENAI_QUOTA_EXHAUSTED", "message": str(exc)}) from exc
    except AiServiceAuthenticationError as exc:
        raise HTTPException(status_code=401, detail={"code": "OPENAI_AUTHENTICATION_FAILED", "message": str(exc)}) from exc
    except AiServiceError as exc:
        raise HTTPException(status_code=503, detail={"code": "AI_FINAL_REPORT_FAILED", "message": str(exc)}) from exc
    except CaseVersionConflictError as exc:
        raise HTTPException(status_code=409, detail={"code": "VERSION_CONFLICT", "message": "Case has changed.", "current_version": exc.current_version}) from exc
    except ValueError as exc:
        raise HTTPException(status_code=409, detail={"code": str(exc), "message": "현재 사건 상태에서는 요청한 변경을 수행할 수 없습니다."}) from exc
    except KeyError as exc:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case not found."}) from exc
    return PublicReportResponse.model_validate(report)


@app.post("/api/cases/{case_id}/reopen", response_model=PublicCaseReadResponse, response_model_exclude_none=True)
async def reopen_closed_case(case_id: str, request: AdminCaseReopenRequest) -> PublicCaseReadResponse:
    require_admin_password(request.password)
    try:
        record = await repository.reopen_case(case_id, request.expected_version)
    except CaseVersionConflictError as exc:
        raise HTTPException(status_code=409, detail={"code": "VERSION_CONFLICT", "message": "Case has changed.", "current_version": exc.current_version}) from exc
    except ValueError as exc:
        raise HTTPException(status_code=409, detail={"code": str(exc), "message": "종결된 사건만 다시 진행할 수 있습니다."}) from exc
    except KeyError as exc:
        raise HTTPException(status_code=404, detail={"code": "CASE_NOT_FOUND", "message": "Case not found."}) from exc
    return await to_case_read(record)


@app.get("/api/cases/{case_id}/reports/final", response_model=PublicReportResponse)
async def get_final_case_report(case_id: str) -> PublicReportResponse:
    await require_case(case_id)
    report = await repository.get_final_report(case_id)
    if report is None:
        raise HTTPException(status_code=404, detail={"code": "FINAL_REPORT_NOT_FOUND", "message": "Final report not found."})
    return PublicReportResponse.model_validate(report)


def _report_export_lines(case_id: str, report: dict) -> list[tuple[str, list[str]]]:
    labels = {
        "executive_summary": "종합 요약", "incident_summary": "사건 개요", "verified_facts": "확인된 사실",
        "actions_taken": "대응 및 처리 내역", "resolution": "최종 해결 결과", "follow_up": "후속 안내", "cautions": "유의사항",
    }
    blocks: list[tuple[str, list[str]]] = [("보고서 정보", [f"Case ID: {case_id}", f"보고서 버전: {report.get('report_version', 1)}", f"생성 시각: {report.get('created_at', '')}"])]
    for section in report.get("sections", []):
        content = section.get("content") or {}
        lines: list[str] = []
        if content.get("text"):
            lines.append(str(content["text"]))
        lines.extend(str(item) for item in content.get("items", []) if str(item).strip())
        if content.get("closure_note"):
            lines.append(f"담당자 종결 메모: {content['closure_note']}")
        if lines:
            blocks.append((labels.get(section.get("section_key"), section.get("section_key", "내용")), lines))
    return blocks


@app.get("/api/cases/{case_id}/reports/final/export")
async def export_final_case_report(case_id: str, format: Literal["pdf", "docx"] = "pdf") -> StreamingResponse:
    await require_case(case_id)
    report = await repository.get_final_report(case_id)
    if report is None:
        raise HTTPException(status_code=404, detail={"code": "FINAL_REPORT_NOT_FOUND", "message": "Final report not found."})
    blocks = _report_export_lines(case_id, report)
    output = io.BytesIO()
    if format == "docx":
        from docx import Document

        document = Document()
        document.add_heading("CSR | Case Share Room 최종 결과 보고서", 0)
        for heading, lines in blocks:
            document.add_heading(heading, level=1)
            for line in lines:
                document.add_paragraph(line, style="List Bullet" if len(lines) > 1 else None)
        document.save(output)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfgen import canvas

        pdfmetrics.registerFont(UnicodeCIDFont("HYSMyeongJo-Medium"))
        page = canvas.Canvas(output)
        page.setTitle(f"CSR {case_id} 최종 결과 보고서")
        width, height = page._pagesize
        y = height - 48
        page.setFont("HYSMyeongJo-Medium", 16)
        page.drawString(44, y, "CSR | Case Share Room 최종 결과 보고서")
        y -= 30
        for heading, lines in blocks:
            if y < 90:
                page.showPage(); y = height - 48
            page.setFont("HYSMyeongJo-Medium", 12); page.drawString(44, y, heading); y -= 20
            page.setFont("HYSMyeongJo-Medium", 9)
            for line in lines:
                chunks = [line[index:index + 64] for index in range(0, len(line), 64)] or [""]
                for chunk in chunks:
                    if y < 58:
                        page.showPage(); page.setFont("HYSMyeongJo-Medium", 9); y = height - 48
                    page.drawString(54, y, f"• {chunk}"); y -= 15
            y -= 8
        page.save()
        media_type = "application/pdf"
    output.seek(0)
    file_name = f"CSR-{case_id}-final-report.{format}"
    return StreamingResponse(output, media_type=media_type, headers={"Content-Disposition": f'attachment; filename="{file_name}"'})
